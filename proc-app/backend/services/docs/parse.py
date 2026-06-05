"""Document parsing — format-routed text extraction.

Each parser returns a list of `Block` records:
  {page: int (1-based), heading: str|None, text: str}

Used by ingest.py which then chunks the blocks into embeddable units.

Formats supported in Phase 1:
  - PDF (text-based)  via pdfplumber
  - DOCX              via python-docx
  - MD / TXT          plain text
  - PNG / JPG         via Gemini vision (single text block describing image)

OCR for scanned PDFs is NOT enabled in Phase 1 (Tesseract dep). When a
PDF returns no extractable text the caller marks status="failed" with a
clear message; consultant can re-upload a text PDF or wait for Phase 2.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

log = logging.getLogger("procvault.docs.parse")


@dataclass
class Block:
    page: int
    heading: Optional[str]
    text: str


class ParseError(Exception):
    """Raised when a document cannot be parsed at all. Caller marks
    status=failed and stores the message."""


def parse_document(path: Path, mime_type: Optional[str] = None) -> tuple[str, list[Block]]:
    """Route to the format-specific parser. Returns (parsed_method, blocks)."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdfplumber", _parse_pdf(path)
    if suffix in (".docx", ".doc"):
        return "docx", _parse_docx(path)
    if suffix in (".md", ".markdown"):
        return "md", _parse_markdown(path)
    if suffix == ".txt":
        return "txt", _parse_plain(path)
    if suffix in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
        return "image", _parse_image(path)
    raise ParseError(f"Unsupported file format: {suffix}")


# ---------------------------------------------------------------------------
# PDF — pdfplumber (text + table flattening, page-aware)
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> list[Block]:
    try:
        import pdfplumber
    except ImportError:
        raise ParseError("pdfplumber not installed")
    blocks: list[Block] = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract text + tables. Tables come back as 2D arrays; we
                # flatten them to pipe-delimited rows so they're searchable.
                text = (page.extract_text() or "").strip()
                tables = page.extract_tables() or []
                table_texts = []
                for t in tables:
                    rows = []
                    for row in t:
                        cells = [str(c or "").strip() for c in row]
                        rows.append(" | ".join(cells))
                    table_texts.append("\n".join(rows))
                if table_texts:
                    text += "\n\n" + "\n\n".join(table_texts)
                if text:
                    heading = _detect_heading(text)
                    blocks.append(Block(page=page_num, heading=heading, text=text))
    except Exception as e:
        raise ParseError(f"pdfplumber failed: {e}") from e
    if not blocks:
        raise ParseError("PDF contained no extractable text (likely scanned). "
                          "Phase 1 doesn't include OCR — re-save as text PDF or "
                          "wait for Phase 2 to enable Tesseract.")
    return blocks


# ---------------------------------------------------------------------------
# DOCX — python-docx (paragraphs grouped under nearest heading)
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> list[Block]:
    try:
        from docx import Document
    except ImportError:
        raise ParseError("python-docx not installed")
    try:
        doc = Document(str(path))
    except Exception as e:
        raise ParseError(f"python-docx failed: {e}") from e
    blocks: list[Block] = []
    current_heading: Optional[str] = None
    buf: list[str] = []
    # Word doesn't have "pages" reliably; we use the heading-block convention
    # and synthesise a sequential page number — 1 per heading group.
    page = 1
    for p in doc.paragraphs:
        style = (p.style.name or "").lower()
        text = p.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            # flush the previous block
            if buf:
                blocks.append(Block(page=page, heading=current_heading, text="\n".join(buf)))
                buf = []
                page += 1
            current_heading = text
        else:
            buf.append(text)
    if buf:
        blocks.append(Block(page=page, heading=current_heading, text="\n".join(buf)))
    # If the doc has no headings, dump everything as one block
    if not blocks:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if all_text:
            blocks.append(Block(page=1, heading=None, text=all_text))
    if not blocks:
        raise ParseError("DOCX contained no extractable text")
    return blocks


# ---------------------------------------------------------------------------
# Markdown — split on top-level headings
# ---------------------------------------------------------------------------

def _parse_markdown(path: Path) -> list[Block]:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks: list[Block] = []
    current_heading: Optional[str] = None
    buf: list[str] = []
    page = 1
    for line in text.splitlines():
        if line.startswith("# ") or line.startswith("## "):
            if buf:
                blocks.append(Block(page=page, heading=current_heading,
                                    text="\n".join(buf).strip()))
                buf = []
                page += 1
            current_heading = line.lstrip("# ").strip()
        else:
            buf.append(line)
    if buf:
        blocks.append(Block(page=page, heading=current_heading,
                            text="\n".join(buf).strip()))
    if not blocks:
        raise ParseError("Empty markdown")
    return blocks


# ---------------------------------------------------------------------------
# Plain text — single block
# ---------------------------------------------------------------------------

def _parse_plain(path: Path) -> list[Block]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        raise ParseError("Empty text file")
    return [Block(page=1, heading=None, text=text)]


# ---------------------------------------------------------------------------
# Image — Gemini vision describes the diagram
# ---------------------------------------------------------------------------

def _parse_image(path: Path) -> list[Block]:
    """Describe the image via Gemini vision. Returns a single block with
    the textual description so it's still chunk+embed compatible."""
    try:
        from .. import llm as llm_service
        from vertexai.generative_models import GenerativeModel, Part
    except Exception:
        raise ParseError("Vertex AI / vision LLM not available — image parsing requires LLM")
    if not llm_service.is_enabled():
        raise ParseError("LLM not configured; image parsing requires Gemini")
    try:
        model = GenerativeModel(llm_service._MODEL_NAME)
        with open(path, "rb") as f:
            img_bytes = f.read()
        mime = {".png": "image/png", ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg", ".gif": "image/gif", ".webp": "image/webp"}.get(path.suffix.lower(), "image/png")
        prompt = ("Describe this procurement document or diagram in detail. "
                  "If it's a process flow, list every step, decision point, "
                  "and approval gate. If it's an org chart, list every role + "
                  "reporting line. If it's a policy excerpt, transcribe key "
                  "rules verbatim. Be specific, list every visible element.")
        resp = model.generate_content([Part.from_data(img_bytes, mime), prompt])
        description = (resp.text or "").strip()
    except Exception as e:
        raise ParseError(f"Image description via Gemini failed: {e}") from e
    if not description:
        raise ParseError("Gemini returned empty description for image")
    return [Block(page=1, heading="(image — Gemini vision description)", text=description)]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _detect_heading(text: str) -> Optional[str]:
    """Best-effort: grab the first short, all-uppercase or Title-Case line
    as the heading. Used for citation context."""
    for line in text.split("\n")[:6]:
        line = line.strip()
        if 5 <= len(line) <= 120:
            words = line.split()
            if (line.isupper() or
                sum(1 for w in words if w[:1].isupper()) >= max(1, len(words) - 1)):
                return line
    return None
